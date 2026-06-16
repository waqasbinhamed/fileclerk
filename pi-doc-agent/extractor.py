"""Text extraction from various document types."""
import json
import logging
from pathlib import Path
from typing import Optional

from rich.console import Console

console = Console()
logger = logging.getLogger(__name__)

# ~1000 tokens ≈ 4000 characters
MAX_CHARS = 4000

SUPPORTED_EXTENSIONS = {
    ".pdf", ".docx", ".xlsx",
    ".py", ".js", ".ts", ".sql", ".sh", ".ipynb",
    ".md", ".txt", ".csv",
}


def extract_text(file_path: "str | Path") -> Optional[dict]:
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix not in SUPPORTED_EXTENSIONS:
        _log_skipped(path, "unsupported type")
        return None

    try:
        if suffix == ".pdf":
            return _extract_pdf(path)
        elif suffix == ".docx":
            return _extract_docx(path)
        elif suffix == ".xlsx":
            return _extract_xlsx(path)
        elif suffix == ".ipynb":
            return _extract_notebook(path)
        else:
            return _extract_plain_text(path)
    except Exception as exc:
        _log_skipped(path, f"extraction error: {exc}")
        logger.warning("Failed to extract %s: %s", path.name, exc)
        return None


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _make_result(path: Path, text: str, method: str) -> dict:
    snippet = _clean_text(text)[:MAX_CHARS]
    return {
        "path": str(path.resolve()),
        "filename": path.name,
        "filetype": path.suffix.lower(),
        "text_snippet": snippet,
        "char_count": len(snippet),
        "extraction_method": method,
    }


def _clean_text(text: str) -> str:
    lines = [line.strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line)


def _extract_pdf(path: Path) -> dict:
    import fitz  # pymupdf

    text = ""
    with fitz.open(str(path)) as doc:
        for page in doc:
            text += page.get_text()
            if len(text) >= MAX_CHARS * 2:
                break

    if len(text.strip()) < 100:
        console.print(
            f"[yellow]PDF text too short (<100 chars), falling back to Tesseract OCR "
            f"(slow on Pi, ~15-30s per page): {path.name}[/yellow]"
        )
        return _extract_pdf_ocr(path)

    return _make_result(path, text, "pymupdf")


def _extract_pdf_ocr(path: Path) -> dict:
    import fitz
    import io
    import pytesseract
    from PIL import Image

    text = ""
    with fitz.open(str(path)) as doc:
        for i, page in enumerate(doc):
            if i >= 3:  # limit OCR pages for Pi performance
                break
            mat = fitz.Matrix(2, 2)  # 2× zoom for better OCR accuracy
            pix = page.get_pixmap(matrix=mat)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            text += pytesseract.image_to_string(img) + "\n"

    return _make_result(path, text, "tesseract_ocr")


def _extract_docx(path: Path) -> dict:
    from docx import Document

    doc = Document(str(path))
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())

    return _make_result(path, "\n".join(parts), "python-docx")


def _extract_xlsx(path: Path) -> dict:
    from openpyxl import load_workbook

    wb = load_workbook(str(path), read_only=True, data_only=True)
    ws = wb.active

    rows = []
    for i, row in enumerate(ws.iter_rows(max_row=20, values_only=True)):
        cells = [str(cell) for cell in row if cell is not None]
        if cells:
            rows.append(" | ".join(cells))

    wb.close()
    return _make_result(path, "\n".join(rows), "openpyxl")


def _extract_notebook(path: Path) -> dict:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        nb = json.load(f)

    sources = []
    for cell in nb.get("cells", []):
        source = cell.get("source", "")
        if isinstance(source, list):
            source = "".join(source)
        if source.strip():
            sources.append(source.strip())

    return _make_result(path, "\n".join(sources), "notebook")


def _extract_plain_text(path: Path) -> dict:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()
    return _make_result(path, text, "plain_text")


def _log_skipped(path: Path, reason: str = "unsupported type") -> None:
    with open("skipped.log", "a") as f:
        f.write(f"{path}\t{reason}\n")
    logger.debug("Skipped %s: %s", path, reason)
