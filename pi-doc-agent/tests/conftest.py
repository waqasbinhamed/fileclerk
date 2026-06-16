"""Shared pytest fixtures for pi-doc-agent tests."""
import json
import sys
from pathlib import Path

import pytest

# Ensure the project root is importable
sys.path.insert(0, str(Path(__file__).parent.parent))


@pytest.fixture
def sample_txt(tmp_path):
    f = tmp_path / "invoice_2024.txt"
    f.write_text(
        "Invoice #1234\nDate: January 15, 2024\nAmount due: $500.00\n"
        "For consulting services rendered.\n" * 20
    )
    return f


@pytest.fixture
def sample_md(tmp_path):
    f = tmp_path / "research_notes.md"
    f.write_text("# Research Paper\n\nThis is a study on machine learning methods.\n\n## Abstract\n\nWe study...\n" * 10)
    return f


@pytest.fixture
def sample_csv(tmp_path):
    f = tmp_path / "expenses.csv"
    f.write_text("name,amount,date\nAlice,100.00,2024-01-15\nBob,200.00,2024-01-20\n" * 10)
    return f


@pytest.fixture
def sample_py(tmp_path):
    f = tmp_path / "script.py"
    f.write_text('def hello():\n    """Greet the user."""\n    print("Hello, world!")\n\nhello()\n')
    return f


@pytest.fixture
def sample_ipynb(tmp_path):
    nb = {
        "nbformat": 4,
        "nbformat_minor": 5,
        "metadata": {},
        "cells": [
            {"cell_type": "code", "source": ["import numpy as np\n", "x = np.array([1, 2, 3])"], "outputs": []},
            {"cell_type": "markdown", "source": ["# Analysis\n", "This notebook analyzes financial data."]},
        ],
    }
    f = tmp_path / "analysis.ipynb"
    f.write_text(json.dumps(nb))
    return f


@pytest.fixture
def sample_docx(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_heading("Invoice for Services", 0)
    doc.add_paragraph("Amount due: $500.00")
    doc.add_paragraph("Date: January 2024")
    path = tmp_path / "invoice.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def sample_xlsx(tmp_path):
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.append(["Item", "Amount", "Date"])
    ws.append(["Service A", 100, "2024-01-01"])
    ws.append(["Service B", 200, "2024-01-15"])
    path = tmp_path / "budget.xlsx"
    wb.save(str(path))
    return path
